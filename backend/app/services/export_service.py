"""
Export service: generates .docx and .pdf reports from analysis results.
"""
import os
import json
import datetime
from typing import Optional

from app.config import settings


def export_analysis_docx(manuscript_title: str, analysis_results: dict) -> str:
    """Generate a .docx report with tracked-changes style analysis."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return _export_fallback_txt(manuscript_title, analysis_results, "docx")

    doc = Document()

    # Title
    title_para = doc.add_heading("Refinery Analysis Report", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Manuscript: {manuscript_title}")
    doc.add_paragraph(f"Generated: {datetime.datetime.utcnow().strftime('%B %d, %Y at %H:%M UTC')}")
    doc.add_paragraph("")

    # Health Dashboard
    revision = analysis_results.get("revision_command", {})
    health = revision.get("health_dashboard", {})
    if health:
        doc.add_heading("Health Dashboard", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Dimension"
        hdr_cells[1].text = "Score"
        for key, value in health.items():
            if key != "overall":
                row = table.add_row().cells
                row[0].text = key.replace("_", " ").title()
                row[1].text = str(value)
        row = table.add_row().cells
        row[0].text = "OVERALL"
        row[1].text = str(health.get("overall", "N/A"))
        doc.add_paragraph("")

    # Module summaries
    module_names = {
        "manuscript_intelligence": "Manuscript Intelligence Engine",
        "voice_isolation": "Voice Isolation Lab",
        "pacing_architect": "Pacing Architect",
        "character_arc": "Character Arc Workshop",
        "prose_refinery": "Prose Refinery",
    }
    for mod_key, mod_name in module_names.items():
        mod_data = analysis_results.get(mod_key, {})
        if not mod_data:
            continue
        doc.add_heading(mod_name, level=1)
        summary = mod_data.get("summary", {})
        for key, value in summary.items():
            doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
        scores = mod_data.get("scores", {})
        if scores:
            doc.add_heading("Scores", level=2)
            for key, value in scores.items():
                doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}/100")
        doc.add_paragraph("")

    # Edit Queue
    edit_queue = revision.get("edit_queue", [])
    if edit_queue:
        doc.add_heading("Revision Edit Queue", level=1)
        doc.add_paragraph(f"Total items: {len(edit_queue)}")
        for item in edit_queue:
            severity_color = {"HIGH": "FF0000", "MEDIUM": "FF8800", "LOW": "0088FF"}
            p = doc.add_paragraph()
            run = p.add_run(f"[{item['severity']}] ")
            run.bold = True
            run = p.add_run(f"{item['finding']}")
            p2 = doc.add_paragraph(f"  Suggestion: {item['suggestion']}")
            p2.style = "List Bullet"
            doc.add_paragraph(f"  Chapter: {item.get('chapter', 'N/A')} | Category: {item.get('category', 'N/A')}")

    # Save
    os.makedirs(settings.EXPORT_DIR, exist_ok=True)
    safe_title = "".join(c for c in manuscript_title if c.isalnum() or c in " -_")[:50]
    filename = f"Refinery_Report_{safe_title}_{datetime.datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(settings.EXPORT_DIR, filename)
    doc.save(filepath)
    return filepath


def export_analysis_pdf(manuscript_title: str, analysis_results: dict) -> str:
    """Generate a PDF report from analysis results."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
    except ImportError:
        return _export_fallback_txt(manuscript_title, analysis_results, "pdf")

    os.makedirs(settings.EXPORT_DIR, exist_ok=True)
    safe_title = "".join(c for c in manuscript_title if c.isalnum() or c in " -_")[:50]
    filename = f"Refinery_Report_{safe_title}_{datetime.datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.pdf"
    filepath = os.path.join(settings.EXPORT_DIR, filename)

    doc = SimpleDocTemplate(filepath, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontSize=24, spaceAfter=20)
    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading1'], fontSize=16, spaceAfter=12, textColor=HexColor('#1a365d'))
    body_style = styles['Normal']

    # Title
    story.append(Paragraph("Refinery Analysis Report", title_style))
    story.append(Paragraph(f"<b>Manuscript:</b> {manuscript_title}", body_style))
    story.append(Paragraph(f"<b>Generated:</b> {datetime.datetime.utcnow().strftime('%B %d, %Y at %H:%M UTC')}", body_style))
    story.append(Spacer(1, 20))

    # Health Dashboard
    revision = analysis_results.get("revision_command", {})
    health = revision.get("health_dashboard", {})
    if health:
        story.append(Paragraph("Health Dashboard", heading_style))
        table_data = [["Dimension", "Score"]]
        for key, value in health.items():
            if key != "overall":
                table_data.append([key.replace("_", " ").title(), str(value)])
        table_data.append(["OVERALL", str(health.get("overall", "N/A"))])
        t = Table(table_data, colWidths=[3.5*inch, 1.5*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#1a365d')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, -1), (-1, -1), HexColor('#e2e8f0')),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('ALIGN', (1, 0), (1, -1), 'CENTER'),
        ]))
        story.append(t)
        story.append(Spacer(1, 15))

    # Module summaries
    module_names = {
        "manuscript_intelligence": "Manuscript Intelligence Engine",
        "voice_isolation": "Voice Isolation Lab",
        "pacing_architect": "Pacing Architect",
        "character_arc": "Character Arc Workshop",
        "prose_refinery": "Prose Refinery",
    }
    for mod_key, mod_name in module_names.items():
        mod_data = analysis_results.get(mod_key, {})
        if not mod_data:
            continue
        story.append(Paragraph(mod_name, heading_style))
        summary = mod_data.get("summary", {})
        for key, value in summary.items():
            story.append(Paragraph(f"<b>{key.replace('_', ' ').title()}:</b> {value}", body_style))
        story.append(Spacer(1, 10))

    # Edit Queue Summary
    edit_queue = revision.get("edit_queue", [])
    if edit_queue:
        story.append(Paragraph("Revision Edit Queue", heading_style))
        for item in edit_queue[:20]:  # Limit to top 20
            severity_colors = {"HIGH": "#e53e3e", "MEDIUM": "#dd6b20", "LOW": "#3182ce"}
            color = severity_colors.get(item["severity"], "#333")
            story.append(Paragraph(
                f'<font color="{color}"><b>[{item["severity"]}]</b></font> {item["finding"]}',
                body_style
            ))
            story.append(Paragraph(f'<i>Suggestion: {item["suggestion"]}</i>', body_style))
            story.append(Spacer(1, 6))

    doc.build(story)
    return filepath


def export_reader_report_pdf(manuscript_title: str, analysis_results: dict) -> str:
    """Generate an Enterprise Reader Report PDF."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
    except ImportError:
        return _export_fallback_txt(manuscript_title, analysis_results, "pdf")

    os.makedirs(settings.EXPORT_DIR, exist_ok=True)
    safe_title = "".join(c for c in manuscript_title if c.isalnum() or c in " -_")[:50]
    filename = f"Reader_Report_{safe_title}_{datetime.datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.pdf"
    filepath = os.path.join(settings.EXPORT_DIR, filename)

    doc = SimpleDocTemplate(filepath, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=22, spaceAfter=15)
    heading_style = ParagraphStyle('Heading', parent=styles['Heading1'], fontSize=14, textColor=HexColor('#1a365d'))

    story.append(Paragraph("Reader Report", title_style))
    story.append(Paragraph(f"<b>Manuscript:</b> {manuscript_title}", styles['Normal']))
    story.append(Paragraph(f"<b>Date:</b> {datetime.datetime.utcnow().strftime('%B %d, %Y')}", styles['Normal']))
    story.append(Spacer(1, 20))

    # Synopsis
    story.append(Paragraph("Synopsis", heading_style))
    mi = analysis_results.get("manuscript_intelligence", {})
    word_count = mi.get("summary", {}).get("total_word_count", "Unknown")
    chapter_count = mi.get("summary", {}).get("chapter_count", "Unknown")
    story.append(Paragraph(
        f"This manuscript spans {word_count:,} words across {chapter_count} chapters. "
        f"It demonstrates a reading level of grade {mi.get('summary', {}).get('reading_level_grade', 'N/A')} "
        f"with an estimated reading time of {mi.get('summary', {}).get('estimated_reading_time_hours', 'N/A')} hours.",
        styles['Normal']
    ))
    story.append(Spacer(1, 10))

    # Acquisition Score
    from app.mock_data.analysis_results import generate_acquisition_score
    acq = generate_acquisition_score(analysis_results)
    story.append(Paragraph("Acquisition Assessment", heading_style))
    story.append(Paragraph(f"<b>Acquisition Score: {acq['acquisition_score']}/100</b>", styles['Normal']))
    story.append(Paragraph(f"<b>Recommendation: {acq['recommendation']}</b>", styles['Normal']))
    story.append(Paragraph(acq['recommendation_detail'], styles['Normal']))
    story.append(Spacer(1, 10))

    # Score breakdown
    story.append(Paragraph("Score Breakdown", heading_style))
    table_data = [["Component", "Score", "Weight", "Weighted"]]
    for comp_name, comp_data in acq["breakdown"].items():
        table_data.append([
            comp_name.replace("_", " ").title(),
            str(comp_data["score"]),
            f"{comp_data['weight']:.0%}",
            str(comp_data["weighted"]),
        ])
    t = Table(table_data, colWidths=[2.5*inch, 1*inch, 1*inch, 1*inch])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), HexColor('#1a365d')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
    ]))
    story.append(t)
    story.append(Spacer(1, 15))

    # Key findings
    story.append(Paragraph("Key Findings", heading_style))
    revision = analysis_results.get("revision_command", {})
    for item in revision.get("edit_queue", [])[:10]:
        story.append(Paragraph(f"<b>[{item['severity']}]</b> {item['finding']}", styles['Normal']))
        story.append(Spacer(1, 4))

    doc.build(story)
    return filepath


def _export_fallback_txt(manuscript_title: str, analysis_results: dict, intended_format: str) -> str:
    """Fallback: export as plain text if libraries are unavailable."""
    os.makedirs(settings.EXPORT_DIR, exist_ok=True)
    safe_title = "".join(c for c in manuscript_title if c.isalnum() or c in " -_")[:50]
    filename = f"Refinery_Report_{safe_title}_{datetime.datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.txt"
    filepath = os.path.join(settings.EXPORT_DIR, filename)

    lines = [
        "=" * 60,
        "REFINERY ANALYSIS REPORT",
        "=" * 60,
        f"Manuscript: {manuscript_title}",
        f"Generated: {datetime.datetime.utcnow().strftime('%B %d, %Y at %H:%M UTC')}",
        f"(Note: {intended_format} export unavailable, using plain text)",
        "",
    ]

    revision = analysis_results.get("revision_command", {})
    health = revision.get("health_dashboard", {})
    if health:
        lines.append("HEALTH DASHBOARD")
        lines.append("-" * 40)
        for key, value in health.items():
            lines.append(f"  {key.replace('_', ' ').title()}: {value}")
        lines.append("")

    for mod_key in ["manuscript_intelligence", "voice_isolation", "pacing_architect", "character_arc", "prose_refinery"]:
        mod_data = analysis_results.get(mod_key, {})
        if not mod_data:
            continue
        lines.append(mod_key.replace("_", " ").upper())
        lines.append("-" * 40)
        for key, value in mod_data.get("summary", {}).items():
            lines.append(f"  {key.replace('_', ' ').title()}: {value}")
        lines.append("")

    edit_queue = revision.get("edit_queue", [])
    if edit_queue:
        lines.append("REVISION EDIT QUEUE")
        lines.append("-" * 40)
        for item in edit_queue:
            lines.append(f"  [{item['severity']}] {item['finding']}")
            lines.append(f"    Suggestion: {item['suggestion']}")
        lines.append("")

    with open(filepath, "w") as f:
        f.write("\n".join(lines))
    return filepath
